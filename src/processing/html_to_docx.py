import json
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_formatting_value(formatting, section, attribute, default):
    return formatting.get(section, {}).get(attribute, default)

def load_formatting_config(config_path):
    with open(config_path, 'r') as file:
        return json.load(file)

def process_html_to_docx(soup, document, formatting):
    # Add title
    title = soup.find('h1')
    if title:
        run = document.add_heading(level=1).add_run(title.text)
        run.font.size = Pt(get_formatting_value(formatting, 'title', 'font_size', 20))
        run.bold = get_formatting_value(formatting, 'title', 'bold', True)
        if 'color' in formatting.get('title', {}):
            run.font.color.rgb = RGBColor.from_string(formatting['title']['color'])
        title.extract()

    # Add contact info
    contact_info = soup.find_all('p')[0:2]
    if contact_info:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for p in contact_info:
            run = contact_paragraph.add_run(p.text + '\n')
            run.font.size = Pt(get_formatting_value(formatting, 'contact_info', 'font_size', 12))
            run.bold = get_formatting_value(formatting, 'contact_info', 'bold', True)
            run.italic = get_formatting_value(formatting, 'contact_info', 'italic', False)
            if 'color' in formatting.get('contact_info', {}):
                run.font.color.rgb = RGBColor.from_string(formatting['contact_info']['color'])
        contact_info[0].extract()
        contact_info[1].extract()

    # Add the rest of the content
    for element in soup.find_all(['h2', 'h3', 'p', 'ul', 'ol']):
        if element.name == 'h2':
            heading = document.add_heading(level=2)
            run = heading.add_run(element.text)
            run.font.size = Pt(get_formatting_value(formatting, 'heading1', 'font_size', 16))
            run.bold = get_formatting_value(formatting, 'heading1', 'bold', True)
            if 'color' in formatting.get('heading1', {}):
                run.font.color.rgb = RGBColor.from_string(formatting['heading1']['color'])
        elif element.name == 'h3':
            heading = document.add_heading(level=3)
            run = heading.add_run(element.text)
            run.font.size = Pt(get_formatting_value(formatting, 'heading2', 'font_size', 14))
            run.bold = get_formatting_value(formatting, 'heading2', 'bold', False)
            run.italic = get_formatting_value(formatting, 'heading2', 'italic', False)
            if 'color' in formatting.get('heading2', {}):
                run.font.color.rgb = RGBColor.from_string(formatting['heading2']['color'])
        elif element.name == 'p':
            p = document.add_paragraph(element.text)
            p.paragraph_format.space_after = Pt(get_formatting_value(formatting, 'paragraph', 'space_after', 8))
            p.style.font.size = Pt(get_formatting_value(formatting, 'paragraph', 'font_size', 12))
            if 'color' in formatting.get('paragraph', {}):
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(formatting['paragraph']['color'])
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = document.add_paragraph(li.text, style='ListBullet')
                p.style.font.size = Pt(get_formatting_value(formatting, 'list_bullet', 'font_size', 12))
                if 'color' in formatting.get('list_bullet', {}):
                    for run in p.runs:
                        run.font.color.rgb = RGBColor.from_string(formatting['list_bullet']['color'])
        elif element.name == 'ol':
            for li in element.find_all('li'):
                p = document.add_paragraph(li.text, style='ListNumber')
                p.style.font.size = Pt(get_formatting_value(formatting, 'list_number', 'font_size', 12))
                if 'color' in formatting.get('list_number', {}):
                    for run in p.runs:
                        run.font.color.rgb = RGBColor.from_string(formatting['list_number']['color'])

def html_to_docx(html, docx_path, config_path):
    document = Document()
    soup = BeautifulSoup(html, 'html.parser')
    formatting = load_formatting_config(config_path)
    process_html_to_docx(soup, document, formatting)
    document.save(docx_path)

if __name__ == "__main__":
    sample_html = """
    <h1>Sample Title</h1>
    <p>Contact Info Line 1</p>
    <p>Contact Info Line 2</p>
    <h2>Sample Heading 1</h2>
    <p>Sample paragraph text.</p>
    <ul>
        <li>Item 1</li>
        <li>Item 2</li>
    </ul>
    """
    html_to_docx(sample_html, "sample_output.docx", "path_to_formatting_config.json")