from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import Tk, filedialog

def read_docx(docx_path):
    document = Document(docx_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def create_formatted_docx(content, save_path):
    document = Document()

    # Define styles
    def add_heading(text, level):
        heading = document.add_heading(level=level)
        run = heading.add_run(text)
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if level > 1 else WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_paragraph(text):
        p = document.add_paragraph(text)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = Pt(1.15)

    # Add title
    add_heading("Ethan Ransberger", 1)

    # Add contact info
    contact_info = document.add_paragraph()
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contact_info.add_run("ethanransberger@gmail.com | Austin, TX | (757) 647-8191\n")
    run.bold = True
    contact_info.add_run("LinkedIn: linkedin.com/in/ethanransberger | GitHub: github.com/EthanRansberger")

    # Process the content
    sections = content.split("---")
    for section in sections:
        lines = section.strip().split("\n")
        if len(lines) > 1:
            heading_text = lines[0].strip()
            content_lines = lines[1:]
            add_heading(heading_text, 2)
            for line in content_lines:
                add_paragraph(line.strip())

    # Save the document
    document.save(save_path)

def select_file(title, filetypes):
    root = Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(title=title, filetypes=filetypes)
    return file_path

def save_file():
    root = Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.asksaveasfilename(title="Save DOCX file", defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    return file_path

# Main script
docx_path = select_file("Select a DOCX file", [("Word files", "*.docx")])
if docx_path:
    content = read_docx(docx_path)
    save_path = save_file()
    if save_path:
        create_formatted_docx(content, save_path)
        print(f"DOCX file saved to {save_path}")
    else:
        print("Save operation cancelled.")
else:
    print("File selection cancelled.")
