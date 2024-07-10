import os
import sys
import json
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import spacy
from fpdf import FPDF

def load_formatting_config(config_path):
    with open(config_path, 'r') as f:
        return json.load(f)

def markdown_to_html(markdown_text):
    return markdown.markdown(markdown_text)

def get_formatting_value(formatting, section, key, default):
    return formatting.get(section, {}).get(key, default)

def html_to_docx(html, docx_path, formatting):
    document = Document()
    soup = BeautifulSoup(html, 'html.parser')

    # Add title
    title = soup.find('h1')
    if title:
        run = document.add_heading(level=1).add_run(title.text)
        run.font.size = Pt(get_formatting_value(formatting, 'title', 'font_size', 20))
        run.bold = get_formatting_value(formatting, 'title', 'bold', True)
        if 'color' in formatting.get('title', {}):
            run.font.color.rgb = RGBColor.from_string(formatting['title']['color'])
        title.extract()

    # Add contact info
    contact_info = soup.find_all('p')[0:2]
    if contact_info:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for p in contact_info:
            run = contact_paragraph.add_run(p.text + '\n')
            run.font.size = Pt(get_formatting_value(formatting, 'contact_info', 'font_size', 12))
            run.bold = get_formatting_value(formatting, 'contact_info', 'bold', True)
            run.italic = get_formatting_value(formatting, 'contact_info', 'italic', False)
            if 'color' in formatting.get('contact_info', {}):
                run.font.color.rgb = RGBColor.from_string(formatting['contact_info']['color'])
        contact_info[0].extract()
        contact_info[1].extract()

    # Add the rest of the content
    for element in soup.find_all(['h2', 'h3', 'p', 'ul', 'ol']):
        if element.name == 'h2':
            heading = document.add_heading(level=2)
            run = heading.add_run(element.text)
            run.font.size = Pt(get_formatting_value(formatting, 'heading1', 'font_size', 16))
            run.bold = get_formatting_value(formatting, 'heading1', 'bold', True)
            if 'color' in formatting.get('heading1', {}):
                run.font.color.rgb = RGBColor.from_string(formatting['heading1']['color'])
        elif element.name == 'h3':
            heading = document.add_heading(level=3)
            run = heading.add_run(element.text)
            run.font.size = Pt(get_formatting_value(formatting, 'heading2', 'font_size', 14))
            run.bold = get_formatting_value(formatting, 'heading2', 'bold', False)
            run.italic = get_formatting_value(formatting, 'heading2', 'italic', False)
            if 'color' in formatting.get('heading2', {}):
                run.font.color.rgb = RGBColor.from_string(formatting['heading2']['color'])
        elif element.name == 'p':
            p = document.add_paragraph(element.text)
            p.paragraph_format.space_after = Pt(get_formatting_value(formatting, 'paragraph', 'space_after', 8))
            p.style.font.size = Pt(get_formatting_value(formatting, 'paragraph', 'font_size', 12))
            if 'color' in formatting.get('paragraph', {}):
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(formatting['paragraph']['color'])
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = document.add_paragraph(li.text, style='ListBullet')
                p.style.font.size = Pt(get_formatting_value(formatting, 'list_bullet', 'font_size', 12))
                if 'color' in formatting.get('list_bullet', {}):
                    for run in p.runs:
                        run.font.color.rgb = RGBColor.from_string(formatting['list_bullet']['color'])
        elif element.name == 'ol':
            for li in element.find_all('li'):
                p = document.add_paragraph(li.text, style='ListNumber')
                p.style.font.size = Pt(get_formatting_value(formatting, 'list_number', 'font_size', 12))
                if 'color' in formatting.get('list_number', {}):
                    for run in p.runs:
                        run.font.color.rgb = RGBColor.from_string(formatting['list_number']['color'])

    document.save(docx_path)

def docx_to_pdf(docx_path, pdf_path):
    doc = Document(docx_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    for para in doc.paragraphs:
        pdf.set_font("Arial", size=12)
        text = para.text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, text)

    pdf.output(pdf_path)
    print(f"PDF file saved to {pdf_path}")

def select_files(title, filetypes):
    if os.environ.get('CI'):
        # If running in a CI environment, use command-line arguments instead of file dialogs
        return os.environ.get(title).split(',')
    else:
        initial_dir = os.path.dirname(os.path.abspath(__file__))  # Set default folder to script's location
        file_paths = sorted([os.path.join(initial_dir, f) for f in os.listdir(initial_dir) if f.endswith(filetypes[0][1][1:])], key=os.path.getmtime)
        if not file_paths:
            raise FileNotFoundError(f"No files found for {filetypes[0][1][1:]} in {initial_dir}")
        return [file_paths[-1]]

def create_output_folders():
    docx_output_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output", "docx")
    pdf_output_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output", "pdf")
    os.makedirs(docx_output_folder, exist_ok=True)
    os.makedirs(pdf_output_folder, exist_ok=True)
    return docx_output_folder, pdf_output_folder

def analyze_resume_with_spacy(text):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    keywords = [token.text for token in doc if token.is_alpha and not token.is_stop]
    print(f"Keywords extracted: {keywords}")

# Main script
if len(sys.argv) > 1 and sys.argv[1] == '--ci':
    # Running in a CI environment
    formatting_config_paths = sys.argv[2].split(',')
    markdown_paths = sys.argv[3].split(',')
else:
    # Running interactively
    formatting_config_paths = select_files("Select the JSON formatting configuration files", [("JSON files", "*.json")])
    markdown_paths = select_files("Select the Markdown files", [("Markdown files", "*.md")])

if markdown_paths and formatting_config_paths:
    docx_output_folder, pdf_output_folder = create_output_folders()
    for markdown_path in markdown_paths:
        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        html_content = markdown_to_html(markdown_text)
        for formatting_config_path in formatting_config_paths:
            formatting = load_formatting_config(formatting_config_path)
            base_filename = os.path.splitext(os.path.basename(markdown_path))[0]
            formatting_name = os.path.splitext(os.path.basename(formatting_config_path))[0]
            docx_filename = f"{base_filename}_{formatting_name}.docx"
            docx_path = os.path.join(docx_output_folder, docx_filename)
            html_to_docx(html_content, docx_path, formatting)
            print(f"DOCX file saved to {docx_path}")
            pdf_filename = f"{base_filename}_{formatting_name}.pdf"
            pdf_path = os.path.join(pdf_output_folder, pdf_filename)
            docx_to_pdf(docx_path, pdf_path)
            analyze_resume_with_spacy(markdown_text)
else:
    print("File selection cancelled.")
